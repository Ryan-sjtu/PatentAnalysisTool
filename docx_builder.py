from docx import Document
from docx.shared import Inches
import datetime
import re
from typing import Dict

def _split_sections(text: str) -> Dict[str, str]:
    """
    尝试把模型输出拆成：brief, full, translation
    不保证100%命中，命中不了就把全文放full。
    """
    t = text.strip()

    # 常见标记尝试（你可按你模型实际输出再加规则）
    brief_patterns = [r"第一篇[:：]\s*简洁版本", r"简洁版本", r"（简洁版）", r"\bBrief\b"]
    full_patterns  = [r"第二篇[:：]\s*完整版本", r"完整版本", r"（完整版）", r"\bFull\b"]
    trans_patterns = [r"原文翻译", r"译文", r"\bTranslation\b"]

    # 找到“译文”起点
    trans_start = None
    for p in trans_patterns:
        m = re.search(p, t)
        if m:
            trans_start = m.start()
            break

    translation = ""
    main = t
    if trans_start is not None:
        main = t[:trans_start].strip()
        translation = t[trans_start:].strip()

    # 再尝试把 main 分成 brief/full
    # 简单策略：如果同时出现“简洁版本”和“完整版本”，就以其位置切
    brief_pos = None
    full_pos = None
    for p in brief_patterns:
        m = re.search(p, main)
        if m:
            brief_pos = m.start()
            break
    for p in full_patterns:
        m = re.search(p, main)
        if m:
            full_pos = m.start()
            break

    brief = ""
    full = ""

    if brief_pos is not None and full_pos is not None and brief_pos < full_pos:
        brief = main[brief_pos:full_pos].strip()
        full = main[full_pos:].strip()
    elif full_pos is not None:
        full = main[full_pos:].strip()
    else:
        full = main

    return {"brief": brief, "full": full, "translation": translation}

def build_docx(out_path: str, pdf_name: str, llm_text: str):
    doc = Document()

    doc.add_heading("专利分析报告（Kimi生成）", level=0)
    doc.add_paragraph(f"来源文件：{pdf_name}")
    doc.add_paragraph(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    parts = _split_sections(llm_text)

    doc.add_heading("第一篇：简洁版本", level=1)
    if parts["brief"].strip():
        for line in parts["brief"].splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("（模型未显式分出简洁版，本次留空或请启用“强制两次调用”策略。）")

    doc.add_page_break()
    doc.add_heading("第二篇：完整版本", level=1)
    for line in parts["full"].splitlines():
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("原文翻译", level=1)
    if parts["translation"].strip():
        for line in parts["translation"].splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("（模型未输出译文段落，如需强制生成可启用二次调用。）")

    doc.save(out_path)
